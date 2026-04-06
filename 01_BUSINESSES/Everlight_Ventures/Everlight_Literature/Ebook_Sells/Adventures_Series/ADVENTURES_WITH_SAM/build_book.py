#!/usr/bin/env python3
"""
build_book.py -- Generate EPUB and KDP-ready DOCX files from MASTER.md manuscripts
for the "Adventures with Sam and Robo" children's book series.

Usage:
    python3 build_book.py 3          # Build Book 3 only
    python3 build_book.py 1 3 5      # Build Books 1, 3, and 5
    python3 build_book.py all        # Build all books (1-5)

Output per book:
    <book_dir>/Sams_<Ordinal>_Superpower.epub
    <book_dir>/Sams_<Ordinal>_Superpower_KDP.docx

Requirements:
    - python-docx (pip install python-docx)
    - Pillow (pip install pillow)
    - Markdown is NOT required -- EPUB is built with zipfile + custom HTML rendering
"""

import os
import re
import sys
import uuid
import zipfile
import textwrap
from pathlib import Path
from datetime import datetime

# ---------------------------------------------------------------------------
# Auto-install missing dependencies
# ---------------------------------------------------------------------------

def ensure_package(import_name, pip_name=None):
    """Try to import a package; install it if missing."""
    pip_name = pip_name or import_name
    try:
        __import__(import_name)
        return True
    except ImportError:
        print(f"  [INFO] '{import_name}' not found. Installing '{pip_name}'...")
        import subprocess
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", pip_name, "-q"],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            print(f"  [ERROR] Failed to install '{pip_name}': {result.stderr.strip()}")
            return False
        try:
            __import__(import_name)
            return True
        except ImportError:
            print(f"  [ERROR] '{import_name}' still not importable after install.")
            return False

HAS_DOCX = ensure_package("docx", "python-docx")
HAS_PIL = ensure_package("PIL", "pillow")

if HAS_DOCX:
    from docx import Document
    from docx.shared import Inches, Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

if HAS_PIL:
    from PIL import Image


# ---------------------------------------------------------------------------
# Series configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent
SERIES_NAME = "Adventures with Sam and Robo"
PUBLISHER = "Everlight Kids | Everlight Ventures"
AUTHOR = "Everlight Kids"
YEAR = str(datetime.now().year)

ORDINALS = {1: "First", 2: "Second", 3: "Third", 4: "Fourth", 5: "Fifth"}

# Book directory mapping -- handles inconsistent naming across the series
BOOK_DIR_MAP = {
    1: "Book1",
    2: "Book 2",
    3: "book_3",
    4: "book_4",
    5: "book_5",
}

# Some books keep MASTER.md in a manuscript/ subdirectory
MASTER_SEARCH_PATTERNS = [
    "Sams_{ordinal}_Superpower_MASTER.md",
    "SAM_BOOK{num}_MASTER.md",
    "manuscript/Sams_{ordinal}_Superpower_MASTER.md",
    "manuscript/SAM_BOOK{num}_MASTER.md",
]


# ---------------------------------------------------------------------------
# Locate book assets
# ---------------------------------------------------------------------------

def find_book_dir(book_num):
    """Return the Path to a book's root directory."""
    dirname = BOOK_DIR_MAP.get(book_num)
    if dirname is None:
        # Try common patterns
        for pattern in [f"Book{book_num}", f"Book_{book_num}", f"book_{book_num}", f"Book {book_num}"]:
            candidate = BASE_DIR / pattern
            if candidate.is_dir():
                return candidate
        return None
    return BASE_DIR / dirname


def find_master_md(book_dir, book_num):
    """Find the MASTER.md file inside a book directory."""
    ordinal = ORDINALS.get(book_num, f"{book_num}th")
    for pat in MASTER_SEARCH_PATTERNS:
        name = pat.format(ordinal=ordinal, num=book_num)
        candidate = book_dir / name
        if candidate.exists():
            return candidate
    # Fallback: glob for any *MASTER*.md
    results = list(book_dir.rglob("*MASTER*.md"))
    # Exclude ILLUSTRATED versions
    results = [r for r in results if "ILLUSTRATED" not in r.name.upper()]
    if results:
        return results[0]
    return None


def find_images_dir(book_dir):
    """Find the images directory for a book."""
    for candidate in [book_dir / "images", book_dir / "Images"]:
        if candidate.is_dir():
            return candidate
    return None


def find_cover_image(images_dir, book_num):
    """Find the cover image (prefer ebook cover, then generic cover)."""
    if images_dir is None:
        return None
    for name in [f"{book_num}_cover_ebook.jpg", f"{book_num}_cover.jpg",
                 f"{book_num}_cover_ebook.png", f"{book_num}_cover.png"]:
        candidate = images_dir / name
        if candidate.exists():
            return candidate
    # Fallback glob
    for ext in ["jpg", "jpeg", "png"]:
        for g in images_dir.glob(f"*cover*.{ext}"):
            if "bw" not in g.name.lower() and "print" not in g.name.lower():
                return g
    return None


# ---------------------------------------------------------------------------
# Parse MASTER.md into structured blocks
# ---------------------------------------------------------------------------

def parse_master_md(filepath):
    """
    Parse a Sam & Robo MASTER.md into structured content.

    Returns dict:
        title: str           -- e.g. "Sam's Third Superpower"
        subtitle: str        -- e.g. "Adventures with Sam & Robo -- Book 3"
        metadata: dict       -- key/value from the header block
        chapters: list       -- [{heading, blocks}]
        back_matter: list    -- [{heading, blocks}]

    Block types:
        paragraph:    {type, text}
        image:        {type, src, alt}
        interactive:  {type, text}
        qa:           {type, question, answer}
        separator:    {type}
        heading3:     {type, text}  -- ### headings in back matter
        table:        {type, rows}  -- markdown table rows
        list_item:    {type, text}
    """
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    result = {
        "title": "",
        "subtitle": "",
        "metadata": {},
        "chapters": [],
        "back_matter": [],
    }

    i = 0

    # --- Parse header ---
    # Line 1: # Title
    while i < len(lines) and not lines[i].strip().startswith("# "):
        i += 1
    if i < len(lines):
        result["title"] = lines[i].strip().lstrip("# ").strip()
        i += 1

    # Line 2: ## Subtitle
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("## ") and "CHAPTER" not in stripped.upper():
            result["subtitle"] = stripped.lstrip("# ").strip()
            i += 1
            break
        elif stripped.startswith("## ") and "CHAPTER" in stripped.upper():
            break  # No subtitle line, jumped to chapter
        i += 1

    # Parse metadata lines (key: value in bold)
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("## "):
            break  # Hit first chapter
        if stripped.startswith("**") and ":**" in stripped:
            # Extract key/value like **Key:** Value
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", stripped)
            if match:
                result["metadata"][match.group(1).strip()] = match.group(2).strip()
        i += 1

    # --- Parse body ---
    current_section = None  # 'chapter' or 'back_matter'
    current_item = None     # current chapter or back matter section
    in_back_matter = False

    while i < len(lines):
        stripped = lines[i].strip()

        # Chapter heading: ## CHAPTER N: Title
        if stripped.startswith("## CHAPTER") or stripped.startswith("## Chapter"):
            in_back_matter = False
            current_item = {"heading": stripped.lstrip("# ").strip(), "blocks": []}
            result["chapters"].append(current_item)
            i += 1
            continue

        # Back matter marker
        if stripped.upper() == "## BACK MATTER":
            in_back_matter = True
            i += 1
            continue

        # ### heading (back matter subsections or in-chapter sub-headings)
        if stripped.startswith("### "):
            heading_text = stripped.lstrip("# ").strip()
            if in_back_matter:
                current_item = {"heading": heading_text, "blocks": []}
                result["back_matter"].append(current_item)
            elif current_item is not None:
                current_item["blocks"].append({"type": "heading3", "text": heading_text})
            i += 1
            continue

        # Image reference: ![alt](path)
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match and current_item is not None:
            current_item["blocks"].append({
                "type": "image",
                "alt": img_match.group(1),
                "src": img_match.group(2),
            })
            i += 1
            continue

        # Separator: ---
        if stripped == "---":
            if current_item is not None:
                current_item["blocks"].append({"type": "separator"})
            i += 1
            continue

        # Interactive Moment
        if stripped.startswith("**Interactive Moment:**"):
            if current_item is not None:
                current_item["blocks"].append({
                    "type": "interactive",
                    "text": stripped.replace("**Interactive Moment:**", "").strip(),
                })
            i += 1
            continue

        # Q&A pair: **Question:** ... / **Answer:** ...
        if stripped.startswith("**Question:**"):
            question = stripped.replace("**Question:**", "").strip()
            answer = ""
            # Look for answer on next non-empty line
            j = i + 1
            while j < len(lines):
                aline = lines[j].strip()
                if aline.startswith("**Answer:**"):
                    answer = aline.replace("**Answer:**", "").strip()
                    j += 1
                    break
                elif aline == "" or aline == "---":
                    break
                j += 1
            if current_item is not None:
                current_item["blocks"].append({
                    "type": "qa",
                    "question": question,
                    "answer": answer,
                })
            i = j
            continue

        # Table rows
        if stripped.startswith("|") and current_item is not None:
            # Collect all consecutive table lines
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                # Skip separator rows like |---|---|
                if not re.match(r"^\|[-\s|:]+\|$", row):
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    table_rows.append(cells)
                i += 1
            current_item["blocks"].append({"type": "table", "rows": table_rows})
            continue

        # List items (- text)
        if stripped.startswith("- ") and current_item is not None:
            current_item["blocks"].append({
                "type": "list_item",
                "text": stripped[2:].strip(),
            })
            i += 1
            continue

        # Numbered list items (1. text)
        num_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_match and current_item is not None:
            current_item["blocks"].append({
                "type": "list_item",
                "text": num_match.group(1).strip(),
            })
            i += 1
            continue

        # Regular paragraph
        if stripped and current_item is not None:
            current_item["blocks"].append({"type": "paragraph", "text": stripped})

        # Skip empty / "End of Master Manuscript" lines
        i += 1

    return result


# ---------------------------------------------------------------------------
# Markdown inline formatting -> HTML
# ---------------------------------------------------------------------------

def md_to_html(text):
    """Convert markdown inline formatting to HTML spans."""
    # Em-dashes
    text = text.replace(" -- ", "\u2014")
    # Bold+italic ***text***
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
    # Bold **text**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic *text*
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


# ===========================================================================
# EPUB BUILDER (pure Python, zipfile-based)
# ===========================================================================

EPUB_CONTAINER_XML = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>"""

EPUB_CSS = """
body {
    font-family: Georgia, "Times New Roman", serif;
    margin: 1em;
    line-height: 1.5;
    color: #222;
}
h1 { font-size: 1.8em; text-align: center; margin-top: 2em; margin-bottom: 0.5em; }
h2 { font-size: 1.4em; text-align: center; margin-top: 1.5em; margin-bottom: 0.3em; }
h3 { font-size: 1.2em; margin-top: 1.2em; margin-bottom: 0.3em; }
p { margin: 0.4em 0; text-indent: 1.5em; }
p.no-indent { text-indent: 0; }
p.center { text-align: center; text-indent: 0; }
.separator { text-align: center; margin: 1em 0; color: #999; }
.interactive {
    background: #f0f8ff; border-left: 4px solid #4a90d9;
    padding: 0.8em; margin: 1em 0; font-style: italic; text-indent: 0;
}
.qa { margin: 0.8em 0; text-indent: 0; }
.qa .question { font-weight: bold; }
.qa .answer { font-style: italic; margin-top: 0.2em; }
img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }
.img-caption { text-align: center; font-size: 0.85em; color: #666; text-indent: 0; }
table { border-collapse: collapse; margin: 1em auto; }
td, th { border: 1px solid #ccc; padding: 0.4em 0.8em; text-align: left; }
th { background: #f5f5f5; font-weight: bold; }
ul, ol { margin: 0.5em 0 0.5em 2em; }
li { margin: 0.2em 0; }
.metadata { font-size: 0.9em; color: #555; text-align: center; margin: 0.3em 0; text-indent: 0; }
.cover-img { max-width: 100%; height: auto; }
"""


def build_epub(parsed, book_num, book_dir, images_dir, cover_path, output_path):
    """Build an EPUB file from parsed MASTER.md content."""
    book_uuid = str(uuid.uuid4())
    ordinal = ORDINALS.get(book_num, f"{book_num}th")
    title = parsed["title"] or f"Sam's {ordinal} Superpower"
    subtitle = parsed["subtitle"] or f"{SERIES_NAME} -- Book {book_num}"

    # Collect all image files referenced in the manuscript
    image_files = {}  # src_path_in_md -> (epub_internal_path, actual_file_path, media_type)
    all_blocks = []
    for ch in parsed["chapters"] + parsed["back_matter"]:
        all_blocks.extend(ch["blocks"])

    for block in all_blocks:
        if block["type"] == "image":
            src = block["src"]
            actual_path = book_dir / src
            if not actual_path.exists():
                # Try images_dir directly
                fname = Path(src).name
                if images_dir:
                    actual_path = images_dir / fname
            if actual_path.exists():
                ext = actual_path.suffix.lower()
                media = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
                epub_path = f"images/{actual_path.name}"
                image_files[src] = (epub_path, actual_path, media)

    # Add cover image
    cover_epub_path = None
    if cover_path and cover_path.exists():
        ext = cover_path.suffix.lower()
        media = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        cover_epub_path = f"images/{cover_path.name}"
        image_files["__cover__"] = (cover_epub_path, cover_path, media)

    # --- Build HTML chapters ---
    html_files = []  # (filename, title, html_content)

    # Cover page
    if cover_epub_path:
        cover_html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>Cover</title><link rel="stylesheet" type="text/css" href="style.css"/></head>
<body>
<div style="text-align:center;">
<img class="cover-img" src="{cover_epub_path}" alt="Cover"/>
</div>
</body>
</html>"""
        html_files.append(("cover.xhtml", "Cover", cover_html))

    # Title page
    title_html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{title}</title><link rel="stylesheet" type="text/css" href="style.css"/></head>
<body>
<h1>{title}</h1>
<p class="center"><i>{subtitle}</i></p>
<p class="center">{PUBLISHER}</p>
<p class="center">&copy; {YEAR} {AUTHOR}</p>
</body>
</html>"""
    html_files.append(("title.xhtml", title, title_html))

    # Chapter pages
    for idx, chapter in enumerate(parsed["chapters"]):
        fname = f"chapter_{idx+1}.xhtml"
        ch_html = _render_blocks_to_html(chapter["heading"], chapter["blocks"], image_files)
        html_files.append((fname, chapter["heading"], ch_html))

    # Back matter (all sections in one file)
    if parsed["back_matter"]:
        bm_parts = []
        for section in parsed["back_matter"]:
            bm_parts.append(f"<h2>{section['heading']}</h2>")
            bm_parts.append(_render_blocks_to_html_body(section["blocks"], image_files))
        bm_html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>Back Matter</title><link rel="stylesheet" type="text/css" href="style.css"/></head>
<body>
{"".join(bm_parts)}
</body>
</html>"""
        html_files.append(("back_matter.xhtml", "Back Matter", bm_html))

    # --- Build content.opf ---
    manifest_items = []
    manifest_items.append('<item id="style" href="style.css" media-type="text/css"/>')
    manifest_items.append('<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>')

    spine_items = []
    for idx, (fname, _, _) in enumerate(html_files):
        item_id = f"html_{idx}"
        manifest_items.append(
            f'<item id="{item_id}" href="{fname}" media-type="application/xhtml+xml"/>'
        )
        spine_items.append(f'<itemref idref="{item_id}"/>')

    for src, (epub_path, actual_path, media) in image_files.items():
        safe_id = re.sub(r"[^a-zA-Z0-9_]", "_", epub_path)
        manifest_items.append(
            f'<item id="{safe_id}" href="{epub_path}" media-type="{media}"/>'
        )

    cover_meta = ""
    if cover_epub_path:
        safe_cover_id = re.sub(r"[^a-zA-Z0-9_]", "_", cover_epub_path)
        cover_meta = f'<meta name="cover" content="{safe_cover_id}"/>'

    content_opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
    <dc:title>{title}</dc:title>
    <dc:creator opf:role="aut">{AUTHOR}</dc:creator>
    <dc:publisher>{PUBLISHER}</dc:publisher>
    <dc:language>en</dc:language>
    <dc:identifier id="BookId">urn:uuid:{book_uuid}</dc:identifier>
    <dc:date>{YEAR}</dc:date>
    {cover_meta}
  </metadata>
  <manifest>
    {chr(10).join(f"    {item}" for item in manifest_items)}
  </manifest>
  <spine toc="ncx">
    {chr(10).join(f"    {item}" for item in spine_items)}
  </spine>
</package>"""

    # --- Build toc.ncx ---
    nav_points = []
    for idx, (fname, ch_title, _) in enumerate(html_files):
        nav_points.append(f"""    <navPoint id="nav_{idx}" playOrder="{idx+1}">
      <navLabel><text>{ch_title}</text></navLabel>
      <content src="{fname}"/>
    </navPoint>""")

    toc_ncx = f"""<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{book_uuid}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{title}</text></docTitle>
  <navMap>
{chr(10).join(nav_points)}
  </navMap>
</ncx>"""

    # --- Write the EPUB zip ---
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        # mimetype must be first and uncompressed
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr("META-INF/container.xml", EPUB_CONTAINER_XML)
        zf.writestr("OEBPS/content.opf", content_opf)
        zf.writestr("OEBPS/toc.ncx", toc_ncx)
        zf.writestr("OEBPS/style.css", EPUB_CSS)

        for fname, _, html_content in html_files:
            zf.writestr(f"OEBPS/{fname}", html_content)

        for src, (epub_path, actual_path, media) in image_files.items():
            zf.write(str(actual_path), f"OEBPS/{epub_path}")

    print(f"    EPUB saved: {output_path}")


def _render_blocks_to_html(heading, blocks, image_files):
    """Render a chapter heading + blocks into a full XHTML document."""
    body = _render_blocks_to_html_body(blocks, image_files)
    return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{heading}</title><link rel="stylesheet" type="text/css" href="style.css"/></head>
<body>
<h1>{heading}</h1>
{body}
</body>
</html>"""


def _render_blocks_to_html_body(blocks, image_files):
    """Render a list of blocks into HTML body content (no wrapper)."""
    parts = []
    first_para = True

    for block in blocks:
        btype = block["type"]

        if btype == "paragraph":
            cls = ' class="no-indent"' if first_para else ""
            parts.append(f"<p{cls}>{md_to_html(block['text'])}</p>")
            first_para = False

        elif btype == "image":
            src = block["src"]
            if src in image_files:
                epub_path = image_files[src][0]
                alt = block.get("alt", "")
                parts.append(f'<p class="center"><img src="{epub_path}" alt="{alt}"/></p>')
            first_para = True

        elif btype == "separator":
            parts.append('<p class="separator">* * *</p>')
            first_para = True

        elif btype == "interactive":
            parts.append(f'<div class="interactive"><b>Interactive Moment:</b> {md_to_html(block["text"])}</div>')
            first_para = True

        elif btype == "qa":
            parts.append(f'<div class="qa"><p class="question">Question: {md_to_html(block["question"])}</p>')
            parts.append(f'<p class="answer">Answer: {md_to_html(block["answer"])}</p></div>')
            first_para = True

        elif btype == "heading3":
            parts.append(f"<h3>{md_to_html(block['text'])}</h3>")
            first_para = True

        elif btype == "table":
            rows = block["rows"]
            parts.append("<table>")
            for ridx, row in enumerate(rows):
                tag = "th" if ridx == 0 else "td"
                cells = "".join(f"<{tag}>{md_to_html(c)}</{tag}>" for c in row)
                parts.append(f"<tr>{cells}</tr>")
            parts.append("</table>")
            first_para = True

        elif btype == "list_item":
            # Accumulate consecutive list items
            parts.append(f"<li>{md_to_html(block['text'])}</li>")
            first_para = True

    # Wrap orphaned <li> in <ul>
    html = "\n".join(parts)
    html = re.sub(r"((?:<li>.*?</li>\n?)+)", r"<ul>\n\1</ul>", html)

    return html


# ===========================================================================
# KDP DOCX BUILDER (using python-docx)
# ===========================================================================

# KDP 8.5x11 specs for illustrated children's books
KDP_PAGE_WIDTH = Inches(8.5)
KDP_PAGE_HEIGHT = Inches(11)
KDP_MARGIN_TOP = Inches(0.75)
KDP_MARGIN_BOTTOM = Inches(0.75)
KDP_MARGIN_INSIDE = Inches(0.875)
KDP_MARGIN_OUTSIDE = Inches(0.5)

BODY_FONT = "Georgia"
BODY_SIZE = Pt(12)
HEADING_SIZE = Pt(20)
SUBHEADING_SIZE = Pt(14)
LINE_SPACING = 1.3


def build_kdp_docx(parsed, book_num, book_dir, images_dir, cover_path, output_path):
    """Build a KDP-ready DOCX from parsed MASTER.md content."""
    if not HAS_DOCX:
        print("    [SKIP] python-docx not available. Cannot build KDP DOCX.")
        print("           Install with: pip install python-docx")
        return

    ordinal = ORDINALS.get(book_num, f"{book_num}th")
    title = parsed["title"] or f"Sam's {ordinal} Superpower"
    subtitle = parsed["subtitle"] or f"{SERIES_NAME} -- Book {book_num}"

    doc = Document()

    # Configure default style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING

    # Set up page size and margins
    section = doc.sections[0]
    section.page_width = KDP_PAGE_WIDTH
    section.page_height = KDP_PAGE_HEIGHT
    section.top_margin = KDP_MARGIN_TOP
    section.bottom_margin = KDP_MARGIN_BOTTOM
    section.left_margin = KDP_MARGIN_INSIDE
    section.right_margin = KDP_MARGIN_OUTSIDE

    # Enable mirror margins
    try:
        sectPr = section._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:mirrorMargins"), "1")
    except Exception:
        pass

    # --- Title page ---
    _add_spacers(doc, 6)
    _add_styled_para(doc, title, font_size=Pt(28), bold=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_para(doc, subtitle, font_size=Pt(16), italic=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
    _add_styled_para(doc, PUBLISHER, font_size=Pt(11), italic=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_para(doc, f"Copyright {YEAR} {AUTHOR}", font_size=Pt(10),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    _add_styled_para(doc, "All rights reserved.", font_size=Pt(10),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Chapters ---
    for chapter in parsed["chapters"]:
        _add_page_break(doc)
        _add_spacers(doc, 2)
        _add_styled_para(doc, chapter["heading"], font_size=HEADING_SIZE, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
        _render_blocks_docx(doc, chapter["blocks"], book_dir, images_dir)

    # --- Back matter ---
    if parsed["back_matter"]:
        _add_page_break(doc)
        _add_styled_para(doc, "BACK MATTER", font_size=Pt(18), bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

        for section_data in parsed["back_matter"]:
            _add_styled_para(doc, section_data["heading"], font_size=SUBHEADING_SIZE,
                             bold=True, space_before=Pt(18), space_after=Pt(12))
            _render_blocks_docx(doc, section_data["blocks"], book_dir, images_dir)

    # --- Save ---
    doc.save(str(output_path))
    print(f"    DOCX saved: {output_path}")


def _render_blocks_docx(doc, blocks, book_dir, images_dir):
    """Render parsed blocks into a python-docx Document."""
    first_para = True

    for block in blocks:
        btype = block["type"]

        if btype == "paragraph":
            indent = Inches(0) if first_para else Inches(0.3)
            _add_rich_para(doc, block["text"], first_line_indent=indent)
            first_para = False

        elif btype == "image":
            img_path = book_dir / block["src"]
            if not img_path.exists() and images_dir:
                img_path = images_dir / Path(block["src"]).name
            if img_path.exists():
                _add_image_to_doc(doc, img_path)
            first_para = True

        elif btype == "separator":
            _add_styled_para(doc, "* * *", font_size=Pt(12),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(12), space_after=Pt(12))
            first_para = True

        elif btype == "interactive":
            p = doc.add_paragraph()
            run_b = p.add_run("Interactive Moment: ")
            run_b.bold = True
            run_b.font.name = BODY_FONT
            run_b.font.size = BODY_SIZE
            _add_inline_runs(p, block["text"])
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            first_para = True

        elif btype == "qa":
            p1 = doc.add_paragraph()
            run_q = p1.add_run("Question: ")
            run_q.bold = True
            run_q.font.name = BODY_FONT
            run_q.font.size = BODY_SIZE
            _add_inline_runs(p1, block["question"])
            p1.paragraph_format.first_line_indent = Inches(0)
            p1.paragraph_format.space_after = Pt(2)

            p2 = doc.add_paragraph()
            run_a = p2.add_run("Answer: ")
            run_a.bold = True
            run_a.font.name = BODY_FONT
            run_a.font.size = BODY_SIZE
            _add_inline_runs(p2, block["answer"])
            p2.paragraph_format.first_line_indent = Inches(0)
            p2.paragraph_format.space_after = Pt(12)
            first_para = True

        elif btype == "heading3":
            _add_styled_para(doc, block["text"], font_size=SUBHEADING_SIZE, bold=True,
                             space_before=Pt(14), space_after=Pt(8))
            first_para = True

        elif btype == "table":
            rows = block["rows"]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ridx, row in enumerate(rows):
                    for cidx, cell_text in enumerate(row):
                        cell = table.cell(ridx, cidx)
                        cell.text = _strip_md(cell_text)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = BODY_FONT
                                run.font.size = Pt(10)
                            if ridx == 0:
                                for run in para.runs:
                                    run.bold = True
            first_para = True

        elif btype == "list_item":
            p = doc.add_paragraph()
            run_bullet = p.add_run("\u2022 ")
            run_bullet.font.name = BODY_FONT
            run_bullet.font.size = BODY_SIZE
            _add_inline_runs(p, block["text"])
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            first_para = True


def _add_image_to_doc(doc, img_path, max_width=Inches(6.5)):
    """Add an image to the DOCX, auto-scaling to fit the page width."""
    try:
        if HAS_PIL:
            with Image.open(img_path) as img:
                w, h = img.size
                dpi = img.info.get("dpi", (150, 150))
                dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
                if dpi_x < 10:
                    dpi_x = 150
                img_width_inches = w / dpi_x
                if img_width_inches > 6.5:
                    width = max_width
                else:
                    width = Inches(img_width_inches)
        else:
            width = Inches(5.0)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
    except Exception as e:
        # If image fails, add placeholder text
        _add_styled_para(doc, f"[Image: {img_path.name}]",
                         font_size=Pt(10), italic=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x99, 0x99, 0x99))


def _add_styled_para(doc, text, font_name=BODY_FONT, font_size=BODY_SIZE,
                     bold=False, italic=False, alignment=None,
                     space_before=None, space_after=None,
                     first_line_indent=None, color=None):
    """Add a simple styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def _add_rich_para(doc, text, first_line_indent=None):
    """Add a paragraph with inline bold/italic formatting."""
    # Replace em-dashes
    text = text.replace(" -- ", "\u2014")

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

    _add_inline_runs(p, text)
    return p


def _add_inline_runs(p, text):
    """Parse markdown bold/italic and add as runs to paragraph p."""
    text = text.replace(" -- ", "\u2014")
    segments = _parse_inline(text)
    for seg_text, is_bold, is_italic in segments:
        run = p.add_run(seg_text)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        run.bold = is_bold
        run.italic = is_italic


def _parse_inline(text):
    """Parse **bold** and *italic* markdown into (text, bold, italic) tuples."""
    segments = []
    # Match ***bold+italic***, **bold**, *italic*
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False))
        if match.group(2) is not None:
            segments.append((match.group(2), True, True))
        elif match.group(3) is not None:
            segments.append((match.group(3), True, False))
        elif match.group(4) is not None:
            segments.append((match.group(4), False, True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    if not segments:
        segments.append((text, False, False))
    return segments


def _strip_md(text):
    """Strip markdown formatting for plain-text contexts."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text.replace(" -- ", "\u2014")


def _add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_spacers(doc, count):
    """Add empty paragraphs for vertical spacing."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


# ===========================================================================
# Main entry point
# ===========================================================================

def build_single_book(book_num):
    """Build EPUB and KDP DOCX for a single book number."""
    ordinal = ORDINALS.get(book_num, f"{book_num}th")
    print(f"\n{'='*60}")
    print(f"  Building Book {book_num}: Sam's {ordinal} Superpower")
    print(f"{'='*60}")

    # Find book directory
    book_dir = find_book_dir(book_num)
    if book_dir is None or not book_dir.is_dir():
        print(f"  [ERROR] Book directory not found for Book {book_num}")
        print(f"          Searched in: {BASE_DIR}")
        return False

    print(f"  Book dir: {book_dir}")

    # Find MASTER.md
    master_md = find_master_md(book_dir, book_num)
    if master_md is None:
        print(f"  [ERROR] No MASTER.md found in {book_dir}")
        return False

    print(f"  MASTER.md: {master_md.name}")

    # Find images
    images_dir = find_images_dir(book_dir)
    if images_dir:
        img_count = len(list(images_dir.glob("*.jpg"))) + len(list(images_dir.glob("*.png")))
        print(f"  Images dir: {images_dir} ({img_count} files)")
    else:
        print(f"  Images dir: (none found)")

    # Find cover
    cover_path = find_cover_image(images_dir, book_num)
    if cover_path:
        print(f"  Cover: {cover_path.name}")
    else:
        print(f"  Cover: (none found)")

    # Parse the manuscript
    parsed = parse_master_md(master_md)
    ch_count = len(parsed["chapters"])
    bm_count = len(parsed["back_matter"])
    print(f"  Parsed: {ch_count} chapters, {bm_count} back matter sections")

    # Determine output directory (same as book_dir for most books,
    # but for books with manuscript/ subdirs, put outputs at book_dir level)
    output_dir = book_dir

    # Build output filenames
    safe_title = f"Sams_{ordinal}_Superpower"
    epub_path = output_dir / f"{safe_title}.epub"
    docx_path = output_dir / f"{safe_title}_KDP.docx"

    # Build EPUB
    print(f"\n  Building EPUB...")
    try:
        build_epub(parsed, book_num, book_dir, images_dir, cover_path, epub_path)
    except Exception as e:
        print(f"    [ERROR] EPUB build failed: {e}")

    # Build KDP DOCX
    print(f"  Building KDP DOCX...")
    try:
        build_kdp_docx(parsed, book_num, book_dir, images_dir, cover_path, docx_path)
    except Exception as e:
        print(f"    [ERROR] DOCX build failed: {e}")

    return True


def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python3 build_book.py 3          # Build Book 3 only")
        print("  python3 build_book.py 1 3 5      # Build Books 1, 3, and 5")
        print("  python3 build_book.py all        # Build all books (1-5)")
        sys.exit(1)

    args = sys.argv[1:]

    if "all" in [a.lower() for a in args]:
        book_nums = [1, 2, 3, 4, 5]
    else:
        book_nums = []
        for a in args:
            try:
                n = int(a)
                if 1 <= n <= 5:
                    book_nums.append(n)
                else:
                    print(f"  [WARN] Book number {n} out of range (1-5), skipping.")
            except ValueError:
                print(f"  [WARN] Invalid argument '{a}', skipping.")

    if not book_nums:
        print("No valid book numbers specified.")
        sys.exit(1)

    print(f"Adventures with Sam and Robo -- Book Builder")
    print(f"Books to build: {', '.join(str(n) for n in book_nums)}")
    print(f"Base directory: {BASE_DIR}")

    success = 0
    for num in book_nums:
        if build_single_book(num):
            success += 1

    print(f"\n{'='*60}")
    print(f"  Done. {success}/{len(book_nums)} books built successfully.")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
