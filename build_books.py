"""
Build KDP-ready DOCX (paperback) and EPUB (ebook) for Sam's Superpower Series.
Books 1, 2, 4, and 5 -- with all images embedded.

Layout rules:
  EBOOK (EPUB + HTML reader): Color images ONLY, no B&W
  PAPERBACK (DOCX): Color images inline with story, B&W coloring pages at end of each chapter
"""
import os
import re
import base64
from io import BytesIO
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ebooklib import epub

# ============================================================
BASE = "/mnt/sdcard/AA_MY_DRIVE/01_BUSINESSES/Everlight_Ventures/Publishing/Ebook_Sells/Adventures_Series/ADVENTURES_WITH_SAM"

BOOKS = [
    {
        "id": 1,
        "title": "Sam's First Superpower",
        "subtitle": "Adventures with Sam and Robo -- Book 1",
        "author": "Everlight Kids",
        "md": f"{BASE}/Book1/Sams_First_Superpower_MASTER.md",
        "img_dir": f"{BASE}/Book1/images",
        "cover": f"{BASE}/Book1/images/1_cover.jpg",
        "out_docx": f"{BASE}/Book1/Sams_First_Superpower_KDP.docx",
        "out_epub": f"{BASE}/Book1/Sams_First_Superpower.epub",
        "reader_html": f"{BASE}/Book1/Sams_First_Superpower_reader.html",
        "prefix": "1",
        "scenes": 12,
    },
    {
        "id": 2,
        "title": "Sam's Second Superpower",
        "subtitle": "Adventures with Sam and Robo -- Book 2",
        "author": "Everlight Kids",
        "md": f"{BASE}/Book 2/Sams_Second_Superpower_MASTER.md",
        "img_dir": f"{BASE}/Book 2/images",
        "cover": f"{BASE}/Book 2/images/2_cover.jpg",
        "out_docx": f"{BASE}/Book 2/Sams_Second_Superpower_KDP.docx",
        "out_epub": f"{BASE}/Book 2/Sams_Second_Superpower.epub",
        "reader_html": f"{BASE}/Book 2/Sams_Second_Superpower_reader.html",
        "prefix": "2",
        "scenes": 11,
    },
    {
        "id": 4,
        "title": "Sam's Fourth Superpower",
        "subtitle": "Adventures with Sam and Robo -- Book 4",
        "author": "Everlight Kids",
        "md": f"{BASE}/book_4/manuscript/Sams_Fourth_Superpower_MASTER.md",
        "img_dir": f"{BASE}/book_4/images",
        "cover": f"{BASE}/book_4/images/4_cover.jpg",
        "out_docx": f"{BASE}/book_4/Sams_Fourth_Superpower_KDP.docx",
        "out_epub": f"{BASE}/book_4/Sams_Fourth_Superpower.epub",
        "reader_html": f"{BASE}/book_4/Sams_Fourth_Superpower_reader.html",
        "prefix": "4",
        "scenes": 12,
    },
    {
        "id": 5,
        "title": "Sam's Fifth Superpower",
        "subtitle": "Adventures with Sam and Robo -- Book 5",
        "author": "Everlight Kids",
        "md": f"{BASE}/book_5/manuscript/Sams_Fifth_Superpower_MASTER.md",
        "img_dir": f"{BASE}/book_5/images",
        "cover": f"{BASE}/book_5/images/5_cover.jpg",
        "out_docx": f"{BASE}/book_5/Sams_Fifth_Superpower_KDP.docx",
        "out_epub": f"{BASE}/book_5/Sams_Fifth_Superpower.epub",
        "reader_html": f"{BASE}/book_5/Sams_Fifth_Superpower_reader.html",
        "prefix": "5",
        "scenes": 12,
    },
]

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
MARGIN_TOP = Inches(0.75)
MARGIN_BOTTOM = Inches(0.75)
MARGIN_INSIDE = Inches(0.75)
MARGIN_OUTSIDE = Inches(0.5)
IMAGE_WIDTH = Inches(4.5)


def compress_image(img_path, max_width=1200, quality=85):
    img = Image.open(img_path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    w, h = img.size
    if w > max_width:
        ratio = max_width / w
        img = img.resize((max_width, int(h * ratio)), Image.LANCZOS)
    buf = BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def parse_md(md_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    blocks = []
    i = 0
    in_back_matter = False

    while i < len(lines):
        line = lines[i].rstrip("\n")
        if not line.strip():
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue

        img_match = re.match(r"!\[([^\]]+)\]\(images/([^\)]+)\)", line.strip())
        if img_match:
            filename = img_match.group(2)
            btype = "image_bw" if "_bw" in filename else "image_color"
            blocks.append({"type": btype, "content": "", "image": filename})
            i += 1
            continue

        if line.startswith("## CHAPTER"):
            blocks.append({"type": "chapter_title", "content": line.lstrip("# ").strip()})
            i += 1
            continue

        if line.startswith("## BACK MATTER"):
            in_back_matter = True
            i += 1
            continue

        if line.startswith("### ") and in_back_matter:
            blocks.append({"type": "section_title", "content": line.lstrip("# ").strip()})
            i += 1
            continue

        if line.startswith("## "):
            blocks.append({"type": "section_title", "content": line.lstrip("# ").strip()})
            i += 1
            continue

        if line.startswith("**Interactive Moment:**"):
            blocks.append({"type": "interactive", "content": line.replace("**Interactive Moment:**", "").strip()})
            i += 1
            continue

        if line.startswith("**Question:**"):
            q = line.replace("**Question:**", "").strip()
            a = ""
            if i + 1 < len(lines) and lines[i + 1].startswith("**Answer:**"):
                a = lines[i + 1].replace("**Answer:**", "").strip()
                i += 1
            blocks.append({"type": "qa", "content": f"Q: {q}\nA: {a}"})
            i += 1
            continue

        if line.startswith("# Sam") or line.startswith("### Everlight"):
            i += 1
            continue
        skip_prefixes = ("**Document Status:", "**Date:", "**Format:", "**Page Layout:",
                         "**Phonics Focus:", "**Core Value:", "**CASEL", "**CCSS",
                         "**Superpower Unlocked:", "**Target Age:", "**Series Position:")
        if any(line.startswith(p) for p in skip_prefixes):
            i += 1
            continue
        if line.startswith("End of Master Manuscript") or line.startswith("Next steps:"):
            i += 1
            continue

        blocks.append({"type": "text", "content": line.strip()})
        i += 1

    return blocks


def strip_md(text):
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text


def add_formatted_paragraph(doc, text, font_size=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(font_size)
        run.font.name = "Georgia"
    return p


# ============================================================
# DOCX -- Color inline, B&W coloring pages at end of chapter
# ============================================================
def build_docx(book):
    print(f"\n  [DOCX] Building {book['title']}...")
    blocks = parse_md(book["md"])
    doc = Document()

    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_INSIDE
    section.right_margin = MARGIN_OUTSIDE

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book["title"])
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Georgia"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book["subtitle"])
    run.font.size = Pt(14)
    run.font.name = "Georgia"
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Everlight Kids | Everlight Ventures")
    run.font.size = Pt(11)
    run.font.name = "Georgia"

    if os.path.exists(book["cover"]):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compressed = compress_image(book["cover"], max_width=1400, quality=90)
        p.add_run().add_picture(BytesIO(compressed), width=IMAGE_WIDTH)

    doc.add_page_break()

    color_count = 0
    bw_count = 0
    pending_bw = []  # collect B&W images for end of chapter

    def flush_bw_pages():
        """Add collected B&W coloring pages at end of current chapter."""
        nonlocal bw_count
        if not pending_bw:
            return

        # Coloring pages section header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("Coloring Pages")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Georgia"
        run.italic = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Color these scenes from the chapter!")
        run.font.size = Pt(11)
        run.font.name = "Georgia"
        run.italic = True

        for bw_path in pending_bw:
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            compressed = compress_image(bw_path, max_width=1400, quality=90)
            p.add_run().add_picture(BytesIO(compressed), width=IMAGE_WIDTH)
            bw_count += 1

        pending_bw.clear()

    for block in blocks:
        btype = block["type"]

        if btype == "chapter_title":
            flush_bw_pages()  # flush previous chapter's B&W pages
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(48)
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run(block["content"])
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = "Georgia"

        elif btype == "section_title":
            flush_bw_pages()  # flush before back matter
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(block["content"])
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Georgia"

        elif btype == "image_color":
            img_path = os.path.join(book["img_dir"], block["image"])
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                compressed = compress_image(img_path, max_width=1400, quality=90)
                p.add_run().add_picture(BytesIO(compressed), width=IMAGE_WIDTH)
                color_count += 1

        elif btype == "image_bw":
            img_path = os.path.join(book["img_dir"], block["image"])
            if os.path.exists(img_path):
                pending_bw.append(img_path)

        elif btype == "interactive":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("Interactive Moment: ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Georgia"
            run = p.add_run(strip_md(block["content"]))
            run.font.size = Pt(11)
            run.font.name = "Georgia"
            run.italic = True

        elif btype == "qa":
            parts = block["content"].split("\n")
            for part in parts:
                p = doc.add_paragraph()
                if part.startswith("Q:"):
                    run = p.add_run("Question: ")
                    run.bold = True
                    run.font.size = Pt(11)
                    run = p.add_run(strip_md(part[2:].strip()))
                    run.font.size = Pt(11)
                elif part.startswith("A:"):
                    run = p.add_run("Answer: ")
                    run.bold = True
                    run.font.size = Pt(11)
                    run = p.add_run(strip_md(part[2:].strip()))
                    run.font.size = Pt(11)

        elif btype == "text":
            content = block["content"]
            if content.startswith("|"):
                p = doc.add_paragraph()
                run = p.add_run(strip_md(content))
                run.font.size = Pt(10)
                run.font.name = "Georgia"
            else:
                add_formatted_paragraph(doc, content)

    flush_bw_pages()  # flush final chapter's B&W pages

    doc.save(book["out_docx"])
    size_mb = os.path.getsize(book["out_docx"]) / (1024 * 1024)
    print(f"  [DOCX] {color_count} color inline + {bw_count} B&W coloring pages, {size_mb:.1f} MB")


# ============================================================
# EPUB -- Color images ONLY, no B&W
# ============================================================
def build_epub(book):
    print(f"\n  [EPUB] Building {book['title']}...")
    blocks = parse_md(book["md"])

    ebook = epub.EpubBook()
    ebook.set_identifier(f"sams-superpower-book-{book['id']}")
    ebook.set_title(book["title"])
    ebook.set_language("en")
    ebook.add_author(book["author"])

    if os.path.exists(book["cover"]):
        cover_data = compress_image(book["cover"], max_width=1600, quality=90)
        ebook.set_cover("cover.jpg", cover_data)

    css = epub.EpubItem(
        uid="style", file_name="style/default.css", media_type="text/css",
        content="""
body { font-family: Georgia, serif; line-height: 1.8; margin: 1em; color: #333; }
h1 { text-align: center; font-size: 1.8em; margin: 2em 0 0.5em; color: #3e2c1a; }
h2 { text-align: center; font-size: 1.4em; margin: 1.5em 0 0.5em; color: #3e2c1a; }
p { margin: 0.5em 0; text-indent: 0; }
.illustration { text-align: center; margin: 1.5em 0; }
.illustration img { max-width: 100%; height: auto; border-radius: 4px; }
.interactive { background: #fff8e1; border: 2px solid #f9a825; border-radius: 8px; padding: 1em; margin: 1.5em 0; }
.interactive-label { font-weight: bold; color: #f57f17; }
.qa { background: #e8f5e9; border-left: 4px solid #43a047; padding: 0.8em 1em; margin: 1em 0; border-radius: 0 8px 8px 0; }
.qa strong { color: #2e7d32; }
.section-title { margin-top: 2em; border-bottom: 2px solid #c4a060; padding-bottom: 0.3em; }
.title-page { text-align: center; margin-top: 30%; }
.title-page h1 { font-size: 2.2em; }
.title-page .subtitle { font-size: 1em; font-style: italic; color: #666; }
.title-page .publisher { font-size: 0.9em; color: #999; margin-top: 2em; }
""".encode("utf-8"),
    )
    ebook.add_item(css)

    # Only add COLOR images to EPUB (no B&W)
    image_items = {}
    for scene in range(1, book["scenes"] + 1):
        fname = f"{book['prefix']}_{scene}_color.jpg"
        fpath = os.path.join(book["img_dir"], fname)
        if os.path.exists(fpath):
            img_data = compress_image(fpath, max_width=1200, quality=85)
            img_item = epub.EpubItem(
                uid=f"img_{scene}_color",
                file_name=f"images/{fname}",
                media_type="image/jpeg",
                content=img_data,
            )
            ebook.add_item(img_item)
            image_items[fname] = f"images/{fname}"

    chapters = []
    current_chapter = None
    current_html = ""
    chapter_count = 0
    img_count = 0

    title_ch = epub.EpubHtml(title="Title", file_name="title.xhtml", lang="en")
    title_ch.content = f"""
    <div class="title-page">
        <h1>{book['title']}</h1>
        <p class="subtitle">{book['subtitle']}</p>
        <p class="publisher">Everlight Kids | Everlight Ventures</p>
    </div>
    """.encode("utf-8")
    title_ch.add_item(css)
    ebook.add_item(title_ch)
    chapters.append(title_ch)

    def flush_chapter():
        nonlocal current_chapter, current_html, chapter_count
        if current_chapter and current_html:
            current_chapter.content = current_html.encode("utf-8")
            current_chapter.add_item(css)
            ebook.add_item(current_chapter)
            chapters.append(current_chapter)

    for block in blocks:
        btype = block["type"]

        if btype == "chapter_title":
            flush_chapter()
            chapter_count += 1
            current_chapter = epub.EpubHtml(
                title=block["content"],
                file_name=f"chapter_{chapter_count}.xhtml",
                lang="en",
            )
            current_html = f'<h1>{block["content"]}</h1>\n'

        elif btype == "section_title":
            if not current_chapter:
                flush_chapter()
                chapter_count += 1
                current_chapter = epub.EpubHtml(
                    title=block["content"],
                    file_name=f"section_{chapter_count}.xhtml",
                    lang="en",
                )
                current_html = ""
            current_html += f'<h2 class="section-title">{block["content"]}</h2>\n'

        elif btype == "image_color":
            fname = block["image"]
            if fname in image_items:
                current_html += f"""
                <div class="illustration">
                    <img src="{image_items[fname]}" alt="Illustration"/>
                </div>\n"""
                img_count += 1

        elif btype == "image_bw":
            pass  # Skip B&W in ebook

        elif btype == "interactive":
            content = strip_md(block["content"])
            current_html += f"""
            <div class="interactive">
                <p><span class="interactive-label">Interactive Moment:</span> {content}</p>
            </div>\n"""

        elif btype == "qa":
            parts = block["content"].split("\n")
            qa_html = '<div class="qa">'
            for part in parts:
                clean = strip_md(part)
                if part.startswith("Q:"):
                    qa_html += f"<p><strong>Question:</strong> {clean[2:].strip()}</p>"
                elif part.startswith("A:"):
                    qa_html += f"<p><strong>Answer:</strong> {clean[2:].strip()}</p>"
            qa_html += "</div>\n"
            current_html += qa_html

        elif btype == "text":
            content = block["content"]
            content = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", content)
            content = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", content)
            if content.startswith("|"):
                current_html += f"<p><small>{content}</small></p>\n"
            else:
                current_html += f"<p>{content}</p>\n"

    flush_chapter()

    ebook.toc = [(ch, []) for ch in chapters]
    ebook.spine = ["nav"] + chapters
    ebook.add_item(epub.EpubNcx())
    ebook.add_item(epub.EpubNav())

    epub.write_epub(book["out_epub"], ebook, {})
    size_mb = os.path.getsize(book["out_epub"]) / (1024 * 1024)
    print(f"  [EPUB] {img_count} color images (no B&W), {size_mb:.1f} MB")


# ============================================================
# HTML READER -- Remove B&W, re-embed color only as base64
# ============================================================
def update_reader_html(book):
    print(f"\n  [HTML] Updating reader for {book['title']}...")
    html_path = book["reader_html"]
    if not os.path.exists(html_path):
        print(f"  [HTML] SKIP -- file not found: {html_path}")
        return

    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    # Remove B&W image paragraphs (base64 or file-linked)
    # Pattern: <p><img src="..." alt="...- B&W Coloring" ...></p>
    bw_pattern = re.compile(
        r'<p><img\s+src="[^"]*"\s+alt="[^"]*B&W[^"]*"[^>]*>\s*</p>',
        re.IGNORECASE
    )
    bw_count = len(bw_pattern.findall(html))
    html = bw_pattern.sub("", html)

    # Also update the toggle label text to remove "B&W" references
    # Change illustration labels from showing two images to just "Full Color Illustration"

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

    size_kb = os.path.getsize(html_path) // 1024
    print(f"  [HTML] Removed {bw_count} B&W images, kept color only. File: {size_kb} KB")


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    for book in BOOKS:
        print(f"\n{'='*60}")
        print(f"  BOOK {book['id']}: {book['title']}")
        print(f"{'='*60}")
        build_epub(book)
        build_docx(book)
        update_reader_html(book)

    print(f"\n{'='*60}")
    print("  ALL DONE!")
    print(f"{'='*60}")
    print("\nOutput files:")
    for book in BOOKS:
        print(f"  Book {book['id']}:")
        print(f"    EPUB (ebook):       {book['out_epub']}")
        print(f"    DOCX (paperback):   {book['out_docx']}")
        print(f"    HTML (phone reader): {book['reader_html']}")
