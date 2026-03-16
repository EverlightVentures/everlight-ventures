#!/usr/bin/env python3
"""
build_kdp_docx.py -- Generate a KDP-ready 6x9 interior .docx file
for "Beyond the Veil: A Quantum Western Thriller" (Hailey Pink Chronicles, Book 1).

KDP Paperback Specs (6x9 inch trim):
  - Margins: Inside 0.875", Outside 0.5", Top 0.75", Bottom 0.75"
  - Body font: Georgia 11pt, line spacing 1.15
  - Chapter titles: 18pt centered, page break before
  - Part titles: 14pt italic centered
  - Scene breaks: centered "* * *"
  - Cipher footers: Courier New, centered, extra spacing above
  - First-line indent: 0.3" (except first para after heading/scene break)
  - Page numbers: centered at bottom, starting on Prologue page

Usage:
    python3 build_kdp_docx.py

Output:
    BEYOND_THE_VEIL_KDP_Interior.docx
"""

import os
import re
import glob
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
CHAPTERS_DIR = BASE_DIR / "chapters"
OUTPUT_FILE = BASE_DIR / "BEYOND_THE_VEIL_KDP_Interior.docx"

CHAPTER_FILES = sorted(CHAPTERS_DIR.glob("*.md"))

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
BOOK_TITLE = "Beyond the Veil"
BOOK_SUBTITLE = "A Quantum Western Thriller"
SERIES_NAME = "The Hailey Pink Chronicles, Book 1"
AUTHOR_NAME = "Hailey Pink"  # pen name placeholder
PUBLISHER = "Everlight Ventures Publishing"
YEAR = "2026"

BODY_FONT = "Georgia"
BODY_SIZE = Pt(11)
HEADING_FONT = "Georgia"
CHAPTER_HEADING_SIZE = Pt(18)
PART_HEADING_SIZE = Pt(14)
CIPHER_FONT = "Courier New"
CIPHER_SIZE = Pt(9)
SCENE_BREAK_TEXT = "*   *   *"

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
MARGIN_INSIDE = Inches(0.875)
MARGIN_OUTSIDE = Inches(0.5)
MARGIN_TOP = Inches(0.75)
MARGIN_BOTTOM = Inches(0.75)

LINE_SPACING = 1.15
FIRST_LINE_INDENT = Inches(0.3)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_section_margins(section, mirror=True):
    """Set page size and margins for a section. mirror=True for gutter-style."""
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    # For mirrored (book) layout: left=inside, right=outside
    section.left_margin = MARGIN_INSIDE
    section.right_margin = MARGIN_OUTSIDE
    section.gutter = Inches(0)
    # Enable mirror margins via XML
    if mirror:
        sectPr = section._sectPr
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            pgMar.set(qn('w:mirrorMargins'), '1')


def set_mirror_margins_globally(doc):
    """Set mirror margins at the document settings level."""
    settings = doc.settings.element
    existing = settings.find(qn('w:mirrorMargins'))
    if existing is None:
        mirror_elem = parse_xml('<w:mirrorMargins {} val="1"/>'.format(nsdecls('w')))
        settings.append(mirror_elem)


def add_page_number_footer(section):
    """Add centered page number to the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear existing runs
    for run in p.runs:
        run.clear()
    p.clear()

    # Add PAGE field
    run1 = p.add_run()
    fldChar1 = parse_xml('<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    run1._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml('<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml('<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run3._element.append(fldChar2)

    # Style the page number
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(10)


def suppress_footer(section):
    """Suppress page numbers / footers for a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()


def add_paragraph_styled(doc, text, font_name=BODY_FONT, font_size=BODY_SIZE,
                         bold=False, italic=False, alignment=None,
                         space_before=None, space_after=None,
                         first_line_indent=None, line_spacing=LINE_SPACING,
                         color=None, keep_with_next=False):
    """Add a paragraph with full styling control."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if keep_with_next:
        pf.keep_with_next = True

    return p


def add_rich_paragraph(doc, text, font_name=BODY_FONT, font_size=BODY_SIZE,
                       alignment=None, first_line_indent=None,
                       space_before=None, space_after=None,
                       line_spacing=LINE_SPACING):
    """Add a paragraph that handles *italic*, **bold**, and em-dashes."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing

    # Replace " -- " with em-dash
    text = text.replace(' -- ', '\u2014')

    # Parse bold and italic markdown
    # Pattern: **bold**, *italic* (but not ** inside **)
    # Process bold first, then italic within each segment
    segments = parse_inline_formatting(text)
    for seg_text, seg_bold, seg_italic in segments:
        run = p.add_run(seg_text)
        run.font.name = font_name
        run.font.size = font_size
        run.bold = seg_bold
        run.italic = seg_italic

    return p


def parse_inline_formatting(text):
    """
    Parse markdown inline formatting: **bold** and *italic*.
    Returns list of (text, is_bold, is_italic) tuples.
    """
    segments = []
    # Regex to find **bold** and *italic* patterns
    # Process bold first (**...**), then italic (*...*)
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False))

        if match.group(2) is not None:
            # Bold: **text**
            segments.append((match.group(2), True, False))
        elif match.group(3) is not None:
            # Italic: *text*
            segments.append((match.group(3), False, True))

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        segments.append((text[pos:], False, False))

    if not segments:
        segments.append((text, False, False))

    return segments


def add_section_break(doc, break_type='new_page'):
    """Add a section break (new page). Returns the new section."""
    from docx.enum.section import WD_SECTION_START
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_margins(new_section)
    return new_section


def add_page_break(doc):
    """Add a simple page break (not a section break)."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ---------------------------------------------------------------------------
# Front Matter
# ---------------------------------------------------------------------------

def build_half_title(doc):
    """Half-title page: just the title, centered vertically via spacing."""
    # Add generous space above to simulate vertical centering
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run()
        run.font.size = Pt(11)

    add_paragraph_styled(
        doc, BOOK_TITLE,
        font_size=Pt(24), bold=False, italic=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(0), line_spacing=1.0
    )


def build_title_page(doc):
    """Full title page with title, subtitle, series, author, publisher."""
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    add_paragraph_styled(
        doc, BOOK_TITLE,
        font_size=Pt(28), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, BOOK_SUBTITLE,
        font_size=Pt(16), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(24), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, SERIES_NAME,
        font_size=Pt(12), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(48), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, "by",
        font_size=Pt(11), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, AUTHOR_NAME,
        font_size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(72), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, PUBLISHER,
        font_size=Pt(10), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(0), line_spacing=1.0
    )


def build_copyright_page(doc):
    """Copyright page."""
    lines = [
        f"Copyright \u00a9 {YEAR} by {AUTHOR_NAME}",
        "All rights reserved.",
        "",
        f"Published by {PUBLISHER}",
        "",
        "This is a work of fiction. Names, characters, places, and incidents "
        "either are the product of the author's imagination or are used "
        "fictitiously. Any resemblance to actual persons, living or dead, "
        "events, or locales is entirely coincidental.",
        "",
        "No part of this book may be reproduced in any form without written "
        "permission from the publisher, except for brief quotations in reviews.",
        "",
        f"First Edition: {YEAR}",
        "",
        "ISBN: [To be assigned]",
        "",
        "Printed in the United States of America",
    ]

    # Some spacing from top
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    for line in lines:
        if line == "":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(6)
        else:
            add_paragraph_styled(
                doc, line,
                font_size=Pt(9),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(2), space_after=Pt(2),
                line_spacing=1.15, first_line_indent=Inches(0)
            )


def build_cipher_key_page(doc):
    """Cipher Key reference page."""
    add_paragraph_styled(
        doc, "The Quantum Web Alphabet",
        font_size=Pt(16), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(36), space_after=Pt(12), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, "Decoded from threads observed in the astral realm",
        font_size=Pt(10), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(18), line_spacing=1.0
    )

    # Flavor text
    flavor_lines = [
        "A hidden voice speaks at the end of each chapter.",
        "Use this key to hear what it says.",
    ]
    for line in flavor_lines:
        add_paragraph_styled(
            doc, line,
            font_size=Pt(10), italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(2), space_after=Pt(2),
            line_spacing=1.15, first_line_indent=Inches(0)
        )

    # Separator
    add_paragraph_styled(
        doc, "",
        font_size=Pt(6),
        space_before=Pt(6), space_after=Pt(6)
    )

    # The cipher key table
    cipher_entries = [
        ("A", "~", "wave pulse"),       ("B", "+", "cross-thread"),
        ("C", "*", "star node"),         ("D", "^", "peak thread"),
        ("E", "#", "lattice point"),     ("F", "@", "vortex eye"),
        ("G", "&", "tangled pair"),      ("H", "%", "split thread"),
        ("I", "!", "singularity"),       ("J", "=", "bridge line"),
        ("K", ";", "pulse pause"),       ("L", ":", "twin nodes"),
        ("M", "/", "ascending thread"),  ("N", "\\", "descending thread"),
        ("O", "(", "open curve"),        ("P", ")", "closed curve"),
        ("Q", "[", "quantum bracket"),   ("R", "]", "mirror bracket"),
        ("S", "{", "spiral open"),       ("T", "}", "spiral close"),
        ("U", "<", "flow inward"),       ("V", ">", "flow outward"),
        ("W", "_", "base thread"),       ("X", "|", "axis line"),
        ("Y", "'", "spark"),             ("Z", "\"", "resonance"),
    ]

    for letter, symbol, name in cipher_entries:
        line = f"{letter}  =  {symbol}        ({name})"
        add_paragraph_styled(
            doc, line,
            font_name=CIPHER_FONT, font_size=Pt(9),
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(0), space_after=Pt(0),
            line_spacing=1.0, first_line_indent=Inches(0.5)
        )

    # Punctuation note
    add_paragraph_styled(
        doc, "",
        font_size=Pt(6),
        space_before=Pt(8), space_after=Pt(4)
    )
    add_paragraph_styled(
        doc, 'Spaces between words are marked with a centered dot:  .',
        font_size=Pt(9), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(4), space_after=Pt(12),
        line_spacing=1.15, first_line_indent=Inches(0)
    )

    # Quick Reference Card
    add_paragraph_styled(
        doc, "Quick Reference",
        font_size=Pt(11), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(6), line_spacing=1.0
    )
    qr_lines = [
        "~  A     +  B     *  C     ^  D     #  E",
        "@  F     &  G     %  H     !  I     =  J",
        ";  K     :  L     /  M     \\  N     (  O",
        ")  P     [  Q     ]  R     {  S     }  T",
        "<  U     >  V     _  W     |  X     '  Y",
        "\"  Z",
    ]
    for line in qr_lines:
        add_paragraph_styled(
            doc, line,
            font_name=CIPHER_FONT, font_size=Pt(9),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(0), space_after=Pt(0),
            line_spacing=1.1, first_line_indent=Inches(0)
        )

    # How to read
    add_paragraph_styled(
        doc, "",
        font_size=Pt(6),
        space_before=Pt(8), space_after=Pt(4)
    )
    instructions = [
        "1. Find the encoded line at the end of each chapter, below the separator.",
        "2. Each symbol represents one letter. Dots represent spaces between words.",
        "3. Work left to right, translating each symbol using the key above.",
        "4. The hidden voice tells a story of its own\u2014one that runs beneath the story you are reading.",
    ]
    for instr in instructions:
        add_paragraph_styled(
            doc, instr,
            font_size=Pt(9),
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(1), space_after=Pt(1),
            line_spacing=1.15, first_line_indent=Inches(0)
        )


def build_dedication_page(doc):
    """Dedication page (Option B)."""
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    dedication_lines = [
        "This book is dedicated to my brother,",
        "who left this world on September 3, 2024.",
        "",
        "You always said the dead do not really leave.",
        "They just go somewhere we cannot follow yet.",
        "",
        "I believe you now.",
        "",
        "This is me finding my way to where you are.",
    ]

    for line in dedication_lines:
        if line == "":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        else:
            add_paragraph_styled(
                doc, line,
                font_size=Pt(12), italic=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(2), space_after=Pt(2),
                line_spacing=1.3, first_line_indent=Inches(0)
            )


def build_toc(doc, chapter_toc_entries):
    """Table of Contents page."""
    add_paragraph_styled(
        doc, "Contents",
        font_size=Pt(18), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(48), space_after=Pt(24), line_spacing=1.0
    )

    for entry in chapter_toc_entries:
        level = entry.get("level", 0)
        title = entry["title"]

        if level == 0:
            # Chapter-level entry
            add_paragraph_styled(
                doc, title,
                font_size=Pt(11), bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(8), space_after=Pt(2),
                line_spacing=1.3,
                first_line_indent=Inches(0)
            )
        else:
            # Part-level entry (indented)
            add_paragraph_styled(
                doc, title,
                font_size=Pt(10), italic=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(1), space_after=Pt(1),
                line_spacing=1.2,
                first_line_indent=Inches(0.4)
            )


# ---------------------------------------------------------------------------
# Chapter Parsing and Rendering
# ---------------------------------------------------------------------------

def parse_chapter_file(filepath):
    """
    Parse a chapter markdown file into structured blocks.

    Returns:
        {
            "chapter_heading": str or None,
            "blocks": [
                {"type": "part_heading", "text": str},
                {"type": "paragraph", "text": str},
                {"type": "scene_break"},
                {"type": "cipher_section", "dots": str, "cipher": str},
                {"type": "end_book_marker"},
            ]
        }
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    result = {
        "chapter_heading": None,
        "blocks": []
    }

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Chapter heading: # Chapter N: Title or # Prologue: Title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            result["chapter_heading"] = stripped[2:].strip()
            i += 1
            continue

        # Part heading: ## Chapter N, Part N: Title
        if stripped.startswith('## '):
            result["blocks"].append({
                "type": "part_heading",
                "text": stripped[3:].strip()
            })
            i += 1
            continue

        # Check for cipher section pattern:
        #   ---
        #   (blank or spaces)
        #   . . .   (with possible leading spaces)
        #   (blank)
        #   <cipher text>
        #   (blank)
        #   ---
        if stripped == '---':
            # Look ahead for cipher pattern: dots line + cipher line
            # Could also be a scene break (just ---) or end-of-chapter separator
            cipher_found = False
            j = i + 1
            # Skip blank lines
            while j < len(lines) and lines[j].strip() == '':
                j += 1
            if j < len(lines) and '. . .' in lines[j]:
                dots_line = lines[j].strip()
                k = j + 1
                # Skip blank lines
                while k < len(lines) and lines[k].strip() == '':
                    k += 1
                if k < len(lines) and lines[k].strip() and not lines[k].strip().startswith('#'):
                    cipher_text = lines[k].strip()
                    # Check if this looks like cipher (contains cipher symbols)
                    if re.match(r'^[~+*^#@&%!=;:/\\()[\]{}|<>_\'".\s-]+$', cipher_text):
                        result["blocks"].append({
                            "type": "cipher_section",
                            "dots": dots_line,
                            "cipher": cipher_text
                        })
                        # Skip past the closing ---
                        m = k + 1
                        while m < len(lines) and lines[m].strip() == '':
                            m += 1
                        if m < len(lines) and lines[m].strip() == '---':
                            i = m + 1
                        else:
                            i = k + 1
                        cipher_found = True

            if cipher_found:
                continue

            # Check for **END OF BOOK 1** pattern after ---
            j2 = i + 1
            while j2 < len(lines) and lines[j2].strip() == '':
                j2 += 1
            if j2 < len(lines) and 'END OF BOOK' in lines[j2].strip().upper():
                result["blocks"].append({"type": "end_book_marker"})
                # Skip past closing ---
                m2 = j2 + 1
                while m2 < len(lines) and lines[m2].strip() == '':
                    m2 += 1
                if m2 < len(lines) and lines[m2].strip() == '---':
                    i = m2 + 1
                else:
                    i = j2 + 1
                continue

            # Otherwise, it is a scene break (---)
            # But only if it is not at the very start (after heading) and not
            # a double --- used as separator
            # Skip multiple consecutive --- lines
            if i > 0:
                # Check if next non-blank line is also ---
                j3 = i + 1
                while j3 < len(lines) and lines[j3].strip() == '':
                    j3 += 1
                if j3 < len(lines) and lines[j3].strip() == '---':
                    # Double ---, skip both (they are decorative separators)
                    i = j3 + 1
                    continue

                # Single --- that is a scene break between paragraphs
                # Only add if there is actual content before and after
                has_content_before = any(
                    b["type"] == "paragraph" for b in result["blocks"]
                )
                if has_content_before:
                    result["blocks"].append({"type": "scene_break"})

            i += 1
            continue

        # Empty line -- skip
        if stripped == '':
            i += 1
            continue

        # Regular paragraph text
        result["blocks"].append({
            "type": "paragraph",
            "text": stripped
        })
        i += 1

    return result


def render_chapter(doc, chapter_data, is_first_chapter=False):
    """Render a parsed chapter into the document."""
    heading = chapter_data["chapter_heading"]
    blocks = chapter_data["blocks"]

    # Page break before chapter (except the very first one which follows
    # the TOC section break)
    if not is_first_chapter:
        add_page_break(doc)

    # Chapter heading: large, centered, with space above
    if heading:
        # Add drop space before title
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        p = add_paragraph_styled(
            doc, heading,
            font_name=HEADING_FONT, font_size=CHAPTER_HEADING_SIZE,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(0), space_after=Pt(24),
            line_spacing=1.0, first_line_indent=Inches(0)
        )

    # Track whether the next paragraph should be un-indented
    # (first para after heading, part heading, or scene break)
    suppress_indent = True

    for block in blocks:
        btype = block["type"]

        if btype == "part_heading":
            # Part heading: centered, italic, 14pt
            add_paragraph_styled(
                doc, "",
                font_size=Pt(6),
                space_before=Pt(6), space_after=Pt(0)
            )
            add_paragraph_styled(
                doc, block["text"],
                font_name=HEADING_FONT, font_size=PART_HEADING_SIZE,
                italic=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(18), space_after=Pt(12),
                line_spacing=1.0, first_line_indent=Inches(0)
            )
            suppress_indent = True

        elif btype == "scene_break":
            add_paragraph_styled(
                doc, SCENE_BREAK_TEXT,
                font_size=Pt(11),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(12), space_after=Pt(12),
                line_spacing=1.0, first_line_indent=Inches(0)
            )
            suppress_indent = True

        elif btype == "paragraph":
            indent = Inches(0) if suppress_indent else FIRST_LINE_INDENT
            add_rich_paragraph(
                doc, block["text"],
                font_name=BODY_FONT, font_size=BODY_SIZE,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=indent,
                space_before=Pt(0), space_after=Pt(2),
                line_spacing=LINE_SPACING
            )
            suppress_indent = False

        elif btype == "cipher_section":
            # Extra space, then dots, then cipher text
            add_paragraph_styled(
                doc, "",
                font_size=Pt(6),
                space_before=Pt(18), space_after=Pt(0)
            )
            add_paragraph_styled(
                doc, block["dots"],
                font_size=Pt(10),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(12), space_after=Pt(6),
                line_spacing=1.0, first_line_indent=Inches(0),
                color=RGBColor(0x66, 0x66, 0x66)
            )
            add_paragraph_styled(
                doc, block["cipher"],
                font_name=CIPHER_FONT, font_size=CIPHER_SIZE,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(4), space_after=Pt(12),
                line_spacing=1.0, first_line_indent=Inches(0)
            )

        elif btype == "end_book_marker":
            add_paragraph_styled(
                doc, "",
                font_size=Pt(6),
                space_before=Pt(18), space_after=Pt(0)
            )
            add_paragraph_styled(
                doc, "END OF BOOK 1",
                font_size=Pt(12), bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(24), space_after=Pt(24),
                line_spacing=1.0, first_line_indent=Inches(0)
            )


# ---------------------------------------------------------------------------
# Back Matter
# ---------------------------------------------------------------------------

def build_about_author(doc):
    """About the Author page."""
    add_page_break(doc)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    add_paragraph_styled(
        doc, "About the Author",
        font_size=Pt(18), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(0), space_after=Pt(24), line_spacing=1.0
    )

    about_text = (
        f"{AUTHOR_NAME} writes fiction that lives in the spaces between genres\u2014"
        "where quantum physics meets the Old West, where grief becomes a doorway, "
        "and where the dead are never as far away as we think. "
        "Beyond the Veil is the first book in The Hailey Pink Chronicles."
    )
    add_paragraph_styled(
        doc, about_text,
        font_size=Pt(11),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12),
        line_spacing=1.3, first_line_indent=Inches(0)
    )

    add_paragraph_styled(
        doc, "[Author bio and photo to be added]",
        font_size=Pt(10), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(24), space_after=Pt(12),
        line_spacing=1.0, first_line_indent=Inches(0),
        color=RGBColor(0x99, 0x99, 0x99)
    )


def build_coming_next(doc):
    """Coming Next / Book 2 teaser page."""
    add_page_break(doc)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    add_paragraph_styled(
        doc, "Coming Next",
        font_size=Pt(18), bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(0), space_after=Pt(18), line_spacing=1.0
    )
    add_paragraph_styled(
        doc, "The Hailey Pink Chronicles, Book 2",
        font_size=Pt(14), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(0), space_after=Pt(36), line_spacing=1.0
    )

    teaser_lines = [
        "Hailey Pink woke in a new timeline.",
        "But the quantum web remembers everything.",
        "",
        "The threads she wove in the astral war did not vanish\u2014they grew.",
        "And something on the other side of the veil noticed.",
        "",
        "The next chapter of the Hailey Pink Chronicles continues the journey",
        "beyond the boundaries of life, death, and the space between.",
        "",
        "Stay tuned.",
    ]

    for line in teaser_lines:
        if line == "":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        else:
            add_paragraph_styled(
                doc, line,
                font_size=Pt(11), italic=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(2), space_after=Pt(2),
                line_spacing=1.3, first_line_indent=Inches(0)
            )


def build_end_matter(doc):
    """Final cipher reveal line."""
    add_page_break(doc)

    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    add_paragraph_styled(
        doc, '"The web remembers every act of love. None of it is lost."',
        font_size=Pt(12), italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(0), space_after=Pt(36),
        line_spacing=1.3, first_line_indent=Inches(0)
    )


# ---------------------------------------------------------------------------
# Build TOC entries from chapter files
# ---------------------------------------------------------------------------

def extract_toc_entries(chapter_files):
    """Extract chapter and part titles from all chapter files for the TOC."""
    entries = []

    for filepath in chapter_files:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                stripped = line.strip()
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    title = stripped[2:].strip()
                    entries.append({"level": 0, "title": title})
                elif stripped.startswith('## '):
                    # Extract just the part title, e.g. "Part 1: Dust and Badges"
                    part_text = stripped[3:].strip()
                    # Try to extract just "Part N: Title" from
                    # "Chapter N, Part N: Title"
                    match = re.match(r'Chapter\s+\d+,\s+(Part\s+\d+:.*)', part_text)
                    if match:
                        entries.append({"level": 1, "title": match.group(1)})
                    else:
                        entries.append({"level": 1, "title": part_text})

    return entries


# ---------------------------------------------------------------------------
# Main Build
# ---------------------------------------------------------------------------

def main():
    print("Building KDP interior .docx ...")
    print(f"  Source chapters: {len(CHAPTER_FILES)} files")
    print(f"  Output: {OUTPUT_FILE}")

    doc = Document()

    # -----------------------------------------------------------------------
    # Configure default styles
    # -----------------------------------------------------------------------
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING

    # -----------------------------------------------------------------------
    # Set up first section (half-title page) -- no page numbers
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    set_section_margins(section)
    set_mirror_margins_globally(doc)
    suppress_footer(section)

    # --- Half-title page ---
    build_half_title(doc)

    # --- Title page (new section, no page numbers) ---
    sec_title = add_section_break(doc)
    suppress_footer(sec_title)
    build_title_page(doc)

    # --- Copyright page (new section, no page numbers) ---
    sec_copy = add_section_break(doc)
    suppress_footer(sec_copy)
    build_copyright_page(doc)

    # --- Cipher Key page (new section, no page numbers) ---
    sec_cipher = add_section_break(doc)
    suppress_footer(sec_cipher)
    build_cipher_key_page(doc)

    # --- Dedication page (new section, no page numbers) ---
    sec_ded = add_section_break(doc)
    suppress_footer(sec_ded)
    build_dedication_page(doc)

    # --- Table of Contents (new section, no page numbers) ---
    sec_toc = add_section_break(doc)
    suppress_footer(sec_toc)
    toc_entries = extract_toc_entries(CHAPTER_FILES)
    build_toc(doc, toc_entries)

    # -----------------------------------------------------------------------
    # Body -- all chapters in one section, with page numbers
    # -----------------------------------------------------------------------
    sec_body = add_section_break(doc)
    add_page_number_footer(sec_body)

    # Reset page numbering to 1 for the body
    sectPr = sec_body._sectPr
    pgNumType = parse_xml('<w:pgNumType {} w:start="1"/>'.format(nsdecls('w')))
    sectPr.append(pgNumType)

    # Render each chapter
    for idx, chapter_file in enumerate(CHAPTER_FILES):
        print(f"  Processing: {chapter_file.name}")
        chapter_data = parse_chapter_file(chapter_file)
        render_chapter(doc, chapter_data, is_first_chapter=(idx == 0))

    # -----------------------------------------------------------------------
    # Back Matter
    # -----------------------------------------------------------------------
    build_about_author(doc)
    build_coming_next(doc)
    build_end_matter(doc)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(str(OUTPUT_FILE))
    print(f"\nDone. File saved to:\n  {OUTPUT_FILE}")
    print(f"\nKDP Interior Specs:")
    print(f"  Trim size: 6\" x 9\"")
    print(f"  Margins: Inside {0.875}\", Outside {0.5}\", Top {0.75}\", Bottom {0.75}\"")
    print(f"  Body font: {BODY_FONT} {11}pt, line spacing {LINE_SPACING}")
    print(f"  Mirror margins: Enabled (for gutter binding)")


if __name__ == "__main__":
    main()
